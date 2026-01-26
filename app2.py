import streamlit as st
from groq import Groq
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re
import random
import datetime
import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth
import bcrypt
import os
import smtplib
from email.mime.text import MIMEText

# ── الإعدادات الفنية ──
st.set_page_config(page_title="مساعد المعلم", layout="wide")

# تحسين التصميم بألوان هادئة (أزرق فاتح، أخضر، رمادي)
st.markdown("""
<style>
    .stApp { background-color: #f0f9ff; }
    .stButton>button { background-color: #3b82f6; color: white; border-radius: 8px; }
    .stButton>button:hover { background-color: #2563eb; }
    h1, h2, h3 { color: #1e40af; text-align: center; }
    .sidebar .sidebar-content { background-color: #d1fae5; padding: 20px; border-radius: 10px; }
    .stTextInput>div>div>input { background-color: #f3f4f6; border: 1px solid #d1d5db; border-radius: 6px; }
    .stExpander { border: 1px solid #e5e7eb; border-radius: 8px; }
</style>
""", unsafe_allow_html=True)

# مفتاح API
MY_API_KEY = "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"

# ── ملف config.yaml لنظام الدخول ──
CONFIG_FILE = 'config.yaml'

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as file:
            return yaml.load(file, Loader=SafeLoader)
    return {
        'credentials': {'usernames': {}},
        'cookie': {'name': 'worksheet_cookie', 'key': 'random_secret_key_123456789', 'expiry_days': 30},
        'pre_authorized': {'emails': []}
    }

def save_config(config):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as file:
        yaml.dump(config, file, allow_unicode=True)

config = load_config()

authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

# ── إضافة شعار رسمي للمنصة في أعلى الصفحة ──
st.image("https://www.moe.gov.sa/ar/PublishingImages/logo.png", width=300, caption="وزارة التعليم - المملكة العربية السعودية")

# ── عرض صفحة تسجيل الدخول ──
authenticator.login()

# ── التحقق من حالة الدخول ──
if st.session_state.get("authentication_status"):
    st.title("مساعد المعلم")

    # ── وضع اسم المعلم تلقائيًا في حقل "اسم المعلم" بعد تسجيل الدخول ──
    if 'teacher_name' not in st.session_state or st.session_state.teacher_name == "":
        st.session_state.teacher_name = st.session_state.get("name", "")

    # ── حفظ القيم الافتراضية في الجلسة (فارغة افتراضيًا) ──
    if 'edu_admin' not in st.session_state:
        st.session_state.edu_admin = ""
    if 'edu_sector' not in st.session_state:
        st.session_state.edu_sector = ""
    if 'school_name' not in st.session_state:
        st.session_state.school_name = ""
    if 'subject_name' not in st.session_state:
        st.session_state.subject_name = "الرياضيات"

    def make_rtl(paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rPr = None
        for child in pPr:
            if child.tag == qn('w:rPr'):
                rPr = child
                break
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            pPr.append(rPr)
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), 'ar-SA')
        lang.set(qn('w:eastAsia'), 'ar-SA')
        lang.set(qn('w:bidi'), 'ar-SA')
        rPr.append(lang)

    def add_page_border(doc, border_size='12'):
        sec = doc.sections[0]._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), border_size)
            border_el.set(qn('w:space'), '24')
            pgBorders.append(border_el)
        sec.append(pgBorders)

    def create_docx(info, questions_data, style, title_color=RGBColor(0, 0, 0), paper_size='A4', orientation='portrait', logo_bytes=None, suggested_time=None):
        doc = Document()
        if style['has_border']: 
            add_page_border(doc, str(style['border_thickness']))
        
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        
        if paper_size == 'A3':
            section.page_width = Inches(11.69)
            section.page_height = Inches(16.54)
        else:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        
        if orientation == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        is_english = info['subject'] == 'اللغة الإنجليزية'

        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.allow_autofit = True

        cells = table.rows[0].cells
        p_logo = cells[1].paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_bytes:
            try:
                doc.add_picture(BytesIO(logo_bytes), width=Inches(2.0))
            except:
                pass

        cells = table.rows[1].cells
        p = cells[0].paragraphs[0]
        p.text = "المملكة العربية السعودية\nوزارة التعليم"
        make_rtl(p)

        p = cells[1].paragraphs[0]
        p.text = f"إدارة التعليم: {info['edu_admin']}\nقطاع التعليم: {info['edu_sector']}\nاسم المدرسة: {info['school']}\nاسم المعلم: {info['teacher']}\nالمادة: {info['subject']}"
        make_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = cells[2].paragraphs[0]
        p.text = f"الصف: {info['grade']}\nالحصة: {info['lesson_period']}\nتاريخ: {datetime.date.today().strftime('%Y-%m-%d')}"
        make_rtl(p)

        p = doc.add_paragraph("✦ ─────────────────────────────── ✦")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(80, 80, 80)

        if suggested_time:
            p = doc.add_paragraph()
            p.text = f"الوقت المقترح للحل: {suggested_time} دقيقة"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(200, 0, 0)
            make_rtl(p)

        p = doc.add_paragraph("\nاسم الطالب: ............................................................")
        make_rtl(p)

        def add_section_title(title_ar, title_en):
            head = doc.add_paragraph()
            title_text = title_en if is_english else title_ar
            run = head.add_run(title_text)
            run.bold = True
            run.font.size = Pt(style['font_size'] + 4)
            run.font.color.rgb = title_color
            make_rtl(head)
            pPr = head._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'F0F0F0')
            pPr.append(shd)
            p_line = doc.add_paragraph("─" * 70)
            p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_line = p_line.runs[0]
            run_line.font.color.rgb = RGBColor(180, 180, 180)
            doc.add_paragraph()

        tf_questions = questions_data['tf'][:]
        mcq_questions = questions_data['mcq'][:]
        essay_questions = questions_data['essay'][:]
        fill_questions = questions_data['fill'][:]

        random.shuffle(tf_questions)
        random.shuffle(mcq_questions)
        random.shuffle(essay_questions)
        random.shuffle(fill_questions)

        for i in range(len(mcq_questions)):
            parts = mcq_questions[i].split("||")
            question_text = parts[0].strip()
            options = [opt.strip() for opt in parts[1:]]
            random.shuffle(options)
            mcq_questions[i] = question_text + " || " + " || ".join(options)

        current_question_index = 0

        if tf_questions:
            add_section_title(
                "Question 1: Mark True or False:" if is_english else "السؤال الأول: ضع علامة (✓) أمام العبارة الصحيحة وعلامة (✗) أمام العبارة الخاطئة:",
                "Question 1: Mark True or False:"
            )
            for i, q in enumerate(tf_questions, 1):
                current_question_index += 1
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = style['line_spacing'] + 0.3
                run_num = p.add_run(f"{i}- ")
                run_num.font.color.rgb = RGBColor(0, 102, 204)
                run_num.font.bold = True
                run_num.font.size = Pt(style['font_size'] + 2)
                run = p.add_run(q)
                run.font.size = Pt(style['font_size'])
                make_rtl(p)
                doc.add_paragraph()

        if mcq_questions:
            doc.add_paragraph()
            add_section_title(
                "Question 2: Choose the correct answer:" if is_english else "السؤال الثاني: اختر الإجابة الصحيحة:",
                "Question 2: Choose the correct answer:"
            )
            for i, q in enumerate(mcq_questions, 1):
                current_question_index += 1
                formatted_q = q.replace("||", "\n")
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = style['line_spacing'] + 0.3
                run_num = p.add_run(f"{i}- ")
                run_num.font.color.rgb = RGBColor(0, 102, 204)
                run_num.font.bold = True
                run_num.font.size = Pt(style['font_size'] + 2)
                run = p.add_run(formatted_q)
                run.font.size = Pt(style['font_size'])
                make_rtl(p)
                doc.add_paragraph()

        if essay_questions:
            doc.add_paragraph()
            add_section_title(
                "Question 3: Answer the following:" if is_english else "السؤال الثالث: أجب عما يأتي:",
                "Question 3: Answer the following:"
            )
            for i, q in enumerate(essay_questions, 1):
                current_question_index += 1
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = style['line_spacing'] + 0.3
                run_num = p.add_run(f"{i}- ")
                run_num.font.color.rgb = RGBColor(0, 102, 204)
                run_num.font.bold = True
                run_num.font.size = Pt(style['font_size'] + 2)
                run = p.add_run(q)
                run.font.size = Pt(style['font_size'])
                make_rtl(p)

                for _ in range(style['essay_lines'] + 1):
                    lp = doc.add_paragraph("................................................................................")
                    make_rtl(lp)
                    lp.paragraph_format.space_after = Pt(8)

        if fill_questions:
            doc.add_paragraph()
            add_section_title(
                "Question 4: Fill in the blanks:" if is_english else "السؤال الرابع: املأ الفراغات بالكلمة المناسبة:",
                "Question 4: Fill in the blanks:"
            )
            for i, q in enumerate(fill_questions, 1):
                current_question_index += 1
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = style['line_spacing'] + 0.3
                run_num = p.add_run(f"{i}- ")
                run_num.font.color.rgb = RGBColor(0, 102, 204)
                run_num.font.bold = True
                run_num.font.size = Pt(style['font_size'] + 2)
                run = p.add_run(q)
                run.font.size = Pt(style['font_size'])
                make_rtl(p)
                doc.add_paragraph()

        paragraph = doc.sections[0].footer.paragraphs[0]
        paragraph.text = "صفحة [PAGE] من [NUMPAGES]"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ── بنك الأسئلة ──
    if 'question_bank' not in st.session_state:
        st.session_state.question_bank = []

    # ── القوائم المنسدلة ──
    SAUDI_GRADES = [
        "الأول الابتدائي", "الثاني الابتدائي", "الثالث الابتدائي", "الرابع الابتدائي", 
        "الخامس الابتدائي", "السادس الابتدائي",
        "الأول المتوسط", "الثاني المتوسط", "الثالث المتوسط",
        "الأول الثانوي", "الثاني الثانوي", "الثالث الثانوي"
    ]

    SUBJECTS = [
        "الرياضيات", "العلوم", "لغتي", "الدراسات الإسلامية", "الاجتماعيات", 
        "اللغة الإنجليزية", "التربية الفنية", "التربية البدنية", "المجالات",
        "العلوم", "اللغة العربية", "الإنجليزية"  # إضافة مواد إضافية
    ]

    DIFFICULTY_LEVELS = ["سهل", "متوسط", "صعب", "تحدي"]

    TITLE_COLORS = {
        "أسود": RGBColor(0, 0, 0),
        "أزرق": RGBColor(0, 0, 255),
        "أحمر": RGBColor(255, 0, 0),
        "أخضر": RGBColor(0, 128, 0)
    }

    # ── الواجهة ──
    st.title("مساعد المعلم")

    if 'tf_list' not in st.session_state: st.session_state.tf_list = []
    if 'mcq_list' not in st.session_state: st.session_state.mcq_list = []
    if 'essay_list' not in st.session_state: st.session_state.essay_list = []
    if 'fill_list' not in st.session_state: st.session_state.fill_list = []
    if 'preview_doc' not in st.session_state: st.session_state.preview_doc = None

    with st.sidebar:
        st.header("📋 البيانات الدراسية")
        edu_admin = st.text_input("إدارة التعليم", value=st.session_state.edu_admin)
        edu_sector = st.text_input("قطاع التعليم", value=st.session_state.edu_sector)
        school = st.text_input("اسم المدرسة", value=st.session_state.school_name)
        teacher = st.text_input("اسم المعلم", value=st.session_state.teacher_name)
        subject = st.selectbox("المادة", SUBJECTS, index=SUBJECTS.index(st.session_state.subject_name) if st.session_state.subject_name in SUBJECTS else 0)
        grade = st.selectbox("الصف", SAUDI_GRADES)
        lesson = st.text_input("عنوان الدرس")
        lesson_period = st.selectbox("الحصة", [f"الحصة {i}" for i in range(1, 8)], index=0)

        st.session_state.edu_admin = edu_admin
        st.session_state.edu_sector = edu_sector
        st.session_state.school_name = school
        st.session_state.teacher_name = teacher
        st.session_state.subject_name = subject

        st.divider()
        st.header("🖼️ شعار المدرسة (اختياري)")
        uploaded_logo = st.file_uploader("ارفع شعار المدرسة (jpg/png)", type=["jpg", "jpeg", "png"])
        logo_bytes = None
        if uploaded_logo is not None:
            try:
                logo_bytes = uploaded_logo.read()
                st.success("تم رفع الشعار بنجاح ✓")
            except:
                st.error("حدث خطأ أثناء رفع الشعار")

        st.divider()
        st.header("🎨 التنسيق والإعدادات")
        f_size = st.slider("حجم الخط", 8, 12, 10)
        l_spacing = st.select_slider("تباعد الأسطر", options=[1.0, 1.15, 1.5], value=1.5)
        border_on = st.checkbox("إطار الصفحة", value=True)
        b_thick = st.slider("سمك الإطار", 4, 20, 12)
        e_lines = st.slider("أسطر المقالي", 1, 4, 2)

        paper_size = st.radio("حجم الورقة", ["A4", "A3"], index=0)
        orientation = st.radio("اتجاه الصفحة", ["عمودي (Portrait)", "أفقي (Landscape)"], index=0)
        title_color_name = st.selectbox("لون عناوين الأسئلة", list(TITLE_COLORS.keys()), index=0)
        title_color = TITLE_COLORS[title_color_name]

        st.divider()
        st.header("⚡ إعدادات متقدمة")
        difficulty = st.selectbox("مستوى الصعوبة", DIFFICULTY_LEVELS, index=1)
        create_two_versions = st.checkbox("إنشاء نسختين مختلفتين (A و B) لمنع الغش", value=False)
        suggested_time = st.number_input("الوقت المقترح للحل (دقائق)", min_value=0, max_value=180, value=45)

    st.subheader("⚙️ هيكل الأسئلة")
    col1, col2, col3, col4 = st.columns(4)
    n_tf   = col1.number_input("عدد (صح/خطأ)"     , 0, 10, 5)
    n_mcq  = col2.number_input("عدد (اختياري)"    , 0, 10, 3)
    n_essay= col3.number_input("عدد (مقالي)"      , 0, 10, 2)
    n_fill = col4.number_input("عدد (ملء فراغات)" , 0, 10, 3)

    # ── مربع نص لإدخال نصوص من خارج الذكاء الاصطناعي ──
    st.header("إدخال نصوص إضافية")
    custom_text = st.text_area("أدخل نصوص إضافية أو تعليمات للورقة (اختياري)")

    if st.button("🚀 إنشاء الورقة"):
        if not lesson:
            st.error("الرجاء كتابة عنوان الدرس")
        else:
            with st.spinner("جاري إعداد الأسئلة وفق المنهج السعودي..."):
                client = Groq(api_key=MY_API_KEY)

                is_english = subject == "اللغة الإنجليزية"
                lang_instruction = "Write the questions in English." if is_english else "Write the questions in Arabic."
                mcq_format_example = "A- Option 1 || B- Option 2 || C- Option 3 || D- Option 4" if is_english else "أ- الخيار الأول || ب- الخيار الثاني || ج- الخيار الثالث || د- الخيار الرابع"

                difficulty_instruction = f"Make the questions at {difficulty} difficulty level." if difficulty != "متوسط" else ""

                prompt = (
                    f"Act as a professional teacher. Create a worksheet for the lesson: '{lesson}' in subject: '{subject}'.\n"
                    f"{lang_instruction}\n"
                    f"{difficulty_instruction}\n"
                    f"Requirements:\n"
                    f"1. Generate {n_tf} True/False sentences. Prefix each with 'TF:'. End with ( ). No answers.\n"
                    f"2. Generate {n_mcq} Multiple Choice Questions. Prefix each with 'MCQ:'. "
                    f"   IMPORTANT FORMAT FOR MCQ: Write the question, then use symbol '||' then write the options separated by '||'.\n"
                    f"   Example: MCQ: Question text? || {mcq_format_example}\n"
                    f"3. Generate {n_essay} Essay questions. Prefix each with 'ESSAY:'. No answers.\n"
                    f"4. Generate {n_fill} Fill-in-the-blanks questions. Prefix each with 'FILL:'. Use _____ for the blank.\n"
                    f"Do NOT include headers like 'Part 1', 'Section A'. Just the prefixed lines."
                )

                try:
                    res = client.chat.completions.create(
                        messages=[{"role":"user","content":prompt}],
                        model="llama-3.3-70b-versatile",
                        temperature=0.4
                    ).choices[0].message.content

                    st.session_state.tf_list = []
                    st.session_state.mcq_list = []
                    st.session_state.essay_list = []
                    st.session_state.fill_list = []

                    lines = res.split('\n')
                    for line in lines:
                        clean_line = line.strip()
                        if clean_line.startswith("TF:"):
                            st.session_state.tf_list.append(clean_line.replace("TF:", "").strip())
                        elif clean_line.startswith("MCQ:"):
                            st.session_state.mcq_list.append(clean_line.replace("MCQ:", "").strip())
                        elif clean_line.startswith("ESSAY:"):
                            st.session_state.essay_list.append(clean_line.replace("ESSAY:", "").strip())
                        elif clean_line.startswith("FILL:"):
                            st.session_state.fill_list.append(clean_line.replace("FILL:", "").strip())

                except Exception as e:
                    st.error(f"حدث خطأ أثناء توليد الأسئلة: {e}")
                    st.stop()

                final_data = {
                    'tf': st.session_state.tf_list,
                    'mcq': st.session_state.mcq_list,
                    'essay': st.session_state.essay_list,
                    'fill': st.session_state.fill_list
                }

                total_questions = len(final_data['tf']) + len(final_data['mcq']) + len(final_data['essay']) + len(final_data['fill'])
                word_count = len(" ".join(final_data['tf'] + final_data['mcq'] + final_data['essay'] + final_data['fill']).split())

                st.session_state.total_questions = total_questions
                st.session_state.word_count = word_count

                style = {
                    'font_size': f_size,
                    'line_spacing': l_spacing,
                    'has_border': border_on,
                    'border_thickness': b_thick,
                    'essay_lines': e_lines
                }

                info = {
                    'school': school,
                    'teacher': teacher,
                    'subject': subject,
                    'grade': grade,
                    'lesson': lesson,
                    'edu_admin': edu_admin,
                    'edu_sector': edu_sector,
                    'lesson_period': lesson_period
                }

                doc_file = create_docx(info, final_data, style, title_color, paper_size, orientation, logo_bytes, suggested_time)
                st.session_state.preview_doc = doc_file

                st.session_state.question_bank.append({
                    'timestamp': datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                    'lesson': lesson,
                    'subject': subject,
                    'grade': grade,
                    'tf': final_data['tf'][:],
                    'mcq': final_data['mcq'][:],
                    'essay': final_data['essay'][:],
                    'fill': final_data['fill'][:]
                })

    if any([st.session_state.tf_list, st.session_state.mcq_list, st.session_state.essay_list, st.session_state.fill_list]):
        st.divider()
        final_data = {'tf': [], 'mcq': [], 'essay': [], 'fill': []}

        if st.session_state.tf_list:
            st.markdown("### ✅ السؤال الأول: (صح / خطأ)")
            for i, q in enumerate(st.session_state.tf_list):
                if st.checkbox(q, key=f"tf_{i}", value=True):
                    final_data['tf'].append(q)

        if st.session_state.mcq_list:
            st.markdown("### 🔢 السؤال الثاني: (اختياري)")
            for i, q in enumerate(st.session_state.mcq_list):
                display_q = q.replace("||", "\n   ")
                if st.checkbox(display_q, key=f"mcq_{i}", value=True):
                    final_data['mcq'].append(q)

        if st.session_state.essay_list:
            st.markdown("### ✍️ السؤال الثالث: (مقالي)")
            for i, q in enumerate(st.session_state.essay_list):
                if st.checkbox(q, key=f"essay_{i}", value=True):
                    final_data['essay'].append(q)

        if st.session_state.fill_list:
            st.markdown("### 🕳️ السؤال الرابع: (ملء الفراغات)")
            for i, q in enumerate(st.session_state.fill_list):
                if st.checkbox(q, key=f"fill_{i}", value=True):
                    final_data['fill'].append(q)

        if any(final_data.values()):
            st.divider()
            style = {
                'font_size': f_size, 
                'line_spacing': l_spacing, 
                'has_border': border_on, 
                'border_thickness': b_thick, 
                'essay_lines': e_lines
            }
            info = {
                'school': school,
                'teacher': teacher,
                'subject': subject,
                'grade': grade,
                'lesson': lesson,
                'edu_admin': edu_admin,
                'edu_sector': edu_sector,
                'lesson_period': lesson_period
            }

            doc_file = create_docx(info, final_data, style, title_color, paper_size, orientation, logo_bytes, suggested_time)

            col1, col2 = st.columns(2)

            with col1:
                st.download_button(
                    "📥 تحميل الورقة للطلاب (Word)",
                    doc_file,
                    file_name=f"{lesson}_للطلاب.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            if st.button("📥 تحميل بنك الأسئلة كملف نصي"):
                bank_text = ""
                for entry in st.session_state.question_bank:
                    bank_text += f"تاريخ: {entry['timestamp']} | المادة: {entry['subject']} | الدرس: {entry['lesson']}\n"
                    bank_text += "صح/خطأ:\n" + "\n".join([f"- {q}" for q in entry['tf']]) + "\n\n"
                    bank_text += "اختياري:\n" + "\n".join([f"- {q}" for q in entry['mcq']]) + "\n\n"
                    bank_text += "مقالي:\n" + "\n".join([f"- {q}" for q in entry['essay']]) + "\n\n"
                    bank_text += "ملء فراغات:\n" + "\n".join([f"- {q}" for q in entry['fill']]) + "\n"
                    bank_text += "─" * 80 + "\n\n"
                
                st.download_button(
                    "تحميل بنك الأسئلة (txt)",
                    bank_text,
                    file_name="بنك_الأسئلة.txt",
                    mime="text/plain"
                )

            if st.button("👁️ معاينة الورقة داخل التطبيق"):
                st.markdown("### معاينة الورقة (نصية مبسطة)")
                st.markdown(f"**{info['school']} - {info['grade']} - {info['subject']}**")
                st.markdown(f"**الدرس: {info['lesson']}**")
                st.markdown("اسم الطالب: .................................................")
                st.markdown("─" * 50)

                if final_data['tf']:
                    st.markdown("**السؤال الأول: صح / خطأ**")
                    for i, q in enumerate(final_data['tf'], 1):
                        st.write(f"{i}- {q}")

                if final_data['mcq']:
                    st.markdown("**السؤال الثاني: اختر الإجابة الصحيحة**")
                    for i, q in enumerate(final_data['mcq'], 1):
                        parts = q.split("||")
                        question = parts[0].strip()
                        options = [opt.strip() for opt in parts[1:]]
                        st.write(f"{i}- {question}")
                        for opt in options:
                            st.write(f"   {opt}")

                if final_data['essay']:
                    st.markdown("**السؤال الثالث: أجب عما يأتي**")
                    for i, q in enumerate(final_data['essay'], 1):
                        st.write(f"{i}- {q}")
                        st.write("   " + "..." * 80)

                if final_data['fill']:
                    st.markdown("**السؤال الرابع: املأ الفراغات**")
                    for i, q in enumerate(final_data['fill'], 1):
                        st.write(f"{i}- {q}")

                st.info("هذه معاينة نصية مبسطة. الملف النهائي (Word) يحتوي على تنسيق أفضل.")

    authenticator.logout("تسجيل الخروج", "sidebar")

elif st.session_state.get("authentication_status") is False:
    st.error("اسم المستخدم أو كلمة المرور غير صحيحة")

elif st.session_state.get("authentication_status") is None:
    # ── حذف العبارة الصفراء (الرجاء إدخال اسم المستخدم وكلمة المرور) ──
    # لا شيء هنا، يظهر فورم الدخول فقط

# ── استعادة كلمة المرور آليًا ──
    if st.button("نسيت كلمة المرور؟"):
     email = st.text_input("أدخل إيميلك")
    if st.button("إرسال كلمة مرور جديدة"):
        if email in config['pre_authorized']['emails']:
            new_password = ''.join(random.choices("abcdefghijklmnopqrstuvwxyz0123456789", k=10))
            hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt(12)).decode('utf-8')
            for username in config['credentials']['usernames']:
                if config['credentials']['usernames'][username].get('email') == email:
                    config['credentials']['usernames'][username]['password'] = hashed_pw
                    save_config(config)
                    msg = MIMEText(f"كلمة المرور الجديدة: {new_password}")
                    msg['Subject'] = 'استعادة كلمة المرور - مساعد المعلم'
                    msg['From'] = 'your_email@gmail.com'
                    msg['To'] = email
                    server = smtplib.SMTP('smtp.gmail.com', 587)
                    server.starttls()
                    server.login("your_email@gmail.com", "your_app_password")
                    server.sendmail("your_email@gmail.com", email, msg.as_string())
                    server.quit()
                    st.success("تم إرسال كلمة مرور جديدة إلى إيميلك!")
                    break
            else:
                st.error("الإيميل غير مسجل")
        else:
            st.error("الإيميل غير مسجل")

# ── تسجيل معلم جديد كقائمة منسدلة (يختفي بعد الدخول) ──
if not st.session_state.get("authentication_status"):
    with st.expander("تسجيل معلم جديد"):
        col1, col2 = st.columns(2)
        with col1:
            new_name = st.text_input("اسم المعلم")
            new_email = st.text_input("الإيميل")
        with col2:
            new_username = st.text_input("اسم المستخدم")
            new_password = st.text_input("كلمة المرور", type="password")

        if st.button("تسجيل المعلم"):
            if new_name and new_email and new_username and new_password:
                # تشفير كلمة المرور
                hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt(12)).decode('utf-8')

                # إضافة المعلم الجديد إلى config
                config['credentials']['usernames'][new_username] = {
                    'name': new_name,
                    'password': hashed_pw,
                    'email': new_email
                }
                config['pre_authorized']['emails'].append(new_email)

                save_config(config)

                st.success(f"تم تسجيل المعلم {new_name} بنجاح! يمكنك الآن تسجيل الدخول باسم المستخدم: {new_username}")
            else:
                st.error("يرجى ملء جميع الحقول")
